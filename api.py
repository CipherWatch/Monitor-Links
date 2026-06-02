import asyncio
import csv
import io
import json
from contextlib import asynccontextmanager

import aiosqlite
from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.requests import Request
from fastapi.responses import JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel

from checker import check_all_urls
from db import DB_PATH, get_all_checks, get_stats, init_db, upsert_link

try:
    import openpyxl
    XLSX_SUPPORT = True
except ImportError:
    XLSX_SUPPORT = False


# ── Config ────────────────────────────────────────────────────────────

def load_config():
    with open("config.json", "r") as f:
        return json.load(f)


def save_config(config):
    with open("config.json", "w") as f:
        json.dump(config, f, indent=2)


# ── Lifespan ──────────────────────────────────────────────────────────

@asynccontextmanager
async def lifespan(app: FastAPI):
    await init_db()
    print("✅ Banco de dados inicializado")
    yield
    print("👋 Servidor encerrado")


# ── App ───────────────────────────────────────────────────────────────

app = FastAPI(
    title="Monitor de Links",
    description="Painel de monitoramento de URLs em tempo real",
    version="1.0.0",
    lifespan=lifespan,
)

app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")


# ── Modelos ───────────────────────────────────────────────────────────

class LinkInput(BaseModel):
    url: str
    name: str
    category: str = "geral"


class SchedulerConfig(BaseModel):
    interval_seconds: int


# ── Rotas HTML ────────────────────────────────────────────────────────

@app.get("/")
async def dashboard(request: Request):
    return templates.TemplateResponse(request=request, name="index.html")


# ── API: Status ───────────────────────────────────────────────────────

@app.get("/api/status")
async def get_status():
    config = load_config()
    results = await check_all_urls(config["urls"], timeout=config["timeout_seconds"])

    data = []
    for item, result in zip(config["urls"], results):
        await upsert_link(item["url"], item["name"], item["category"])
        stats = await get_stats(item["url"])
        data.append({
            "name": result.name,
            "url": result.url,
            "category": item.get("category", "geral"),
            "status": result.status.value,
            "status_code": result.status_code,
            "response_time_ms": result.response_time_ms,
            "error_message": result.error_message,
            "uptime_pct": stats.get("uptime_pct") if stats else None,
            "avg_ms": stats.get("avg_ms") if stats else None,
            "total_checks": stats.get("total") if stats else 0,
        })

    return JSONResponse(content={"results": data})


# ── API: Histórico ────────────────────────────────────────────────────

@app.get("/api/history")
async def get_history_api(url: str | None = None, limit: int = 50):
    rows = await get_all_checks(limit=limit)
    if url:
        rows = [r for r in rows if r["url"] == url]
    return JSONResponse(content={"history": rows})


# ── API: Alertas ──────────────────────────────────────────────────────

@app.get("/api/alerts")
async def get_alerts_api(limit: int = 20):
    async with aiosqlite.connect(DB_PATH) as db:
        db.row_factory = aiosqlite.Row
        async with db.execute(
            "SELECT * FROM alerts ORDER BY triggered_at DESC LIMIT ?", (limit,)
        ) as cur:
            rows = [dict(r) for r in await cur.fetchall()]
    return JSONResponse(content={"alerts": rows})


# ── API: Links ────────────────────────────────────────────────────────

@app.get("/api/links")
async def list_links():
    async with aiosqlite.connect(DB_PATH) as db:
        db.row_factory = aiosqlite.Row
        async with db.execute("SELECT * FROM links ORDER BY name") as cur:
            rows = [dict(r) for r in await cur.fetchall()]
    return JSONResponse(content={"links": rows})


@app.post("/api/links")
async def add_link(link: LinkInput):
    config = load_config()
    if any(u["url"] == link.url for u in config["urls"]):
        raise HTTPException(status_code=409, detail="URL já cadastrada")
    config["urls"].append({"url": link.url, "name": link.name, "category": link.category})
    save_config(config)
    await upsert_link(link.url, link.name, link.category)
    return {"message": "URL adicionada", "link": link.model_dump()}


@app.delete("/api/links")
async def remove_link(url: str):
    config = load_config()
    original = len(config["urls"])
    config["urls"] = [u for u in config["urls"] if u["url"] != url]
    if len(config["urls"]) == original:
        raise HTTPException(status_code=404, detail="URL não encontrada")
    save_config(config)
    return {"message": "URL removida"}


# ── API: Importar planilha ────────────────────────────────────────────

@app.post("/api/links/import")
async def import_links(file: UploadFile = File(...)):
    content  = await file.read()
    filename = file.filename.lower()
    rows     = []

    if filename.endswith(".csv"):
        text   = content.decode("utf-8-sig")
        reader = csv.DictReader(io.StringIO(text))
        rows   = list(reader)

    elif filename.endswith(".xlsx"):
        if not XLSX_SUPPORT:
            return JSONResponse(status_code=400, content={"error": "openpyxl não instalado"})
        wb      = openpyxl.load_workbook(io.BytesIO(content))
        ws      = wb.active
        headers = [str(c.value).strip().lower() for c in next(ws.iter_rows(min_row=1, max_row=1))]
        for row in ws.iter_rows(min_row=2, values_only=True):
            rows.append(dict(zip(headers, row)))
    elif filename.endswith(".docx"):
        try:
            from docx import Document
            from urllib.parse import urlparse
            import re
            doc = Document(io.BytesIO(content))
            for table in doc.tables:
                headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
                for row in table.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    row_dict = dict(zip(headers, cells))
                    url  = row_dict.get("url") or row_dict.get("link") or ""
                    name = row_dict.get("name") or row_dict.get("nome") or ""
                    if url:
                        rows.append({"url": url, "name": name})
            if not rows:
                url_pattern = re.compile(r"https?://[^\s]+")
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    for u in url_pattern.findall(text):
                        name_part = text.replace(u, "").strip(" -|:")
                        rows.append({"url": u, "name": name_part or urlparse(u).netloc})
        except Exception as e:
            return JSONResponse(status_code=400, content={"error": f"Erro ao ler .docx: {str(e)}"})
    else:
        return JSONResponse(status_code=400, content={"error": "Use .csv, .xlsx ou .docx"})

    config        = load_config()
    existing_urls = {u["url"] for u in config["urls"]}
    imported, skipped, errors = [], [], []

    for i, row in enumerate(rows, start=2):
        row_lower = {k.strip().lower(): str(v).strip() for k, v in row.items() if v}
        url  = row_lower.get("url") or row_lower.get("link") or ""
        name = row_lower.get("name") or row_lower.get("nome") or ""

        if not url:
            errors.append(f"Linha {i}: URL vazia")
            continue
        if not url.startswith("http"):
            errors.append(f"Linha {i}: URL inválida ({url})")
            continue
        if not name:
            from urllib.parse import urlparse
            name = urlparse(url).netloc or url
        if url in existing_urls:
            skipped.append(name)
            continue

        config["urls"].append({"url": url, "name": name, "category": "importado"})
        existing_urls.add(url)
        await upsert_link(url, name, "importado")
        imported.append(name)

    save_config(config)
    return JSONResponse(content={"imported": imported, "skipped": skipped, "errors": errors, "total": len(imported)})


# ── API: Scheduler ────────────────────────────────────────────────────

@app.get("/api/scheduler")
async def get_scheduler():
    config = load_config()
    return {"interval_seconds": config.get("check_interval_seconds", 60)}


@app.post("/api/scheduler")
async def set_scheduler(cfg: SchedulerConfig):
    if cfg.interval_seconds < 10:
        raise HTTPException(status_code=400, detail="Intervalo mínimo é 10 segundos")
    config = load_config()
    config["check_interval_seconds"] = cfg.interval_seconds
    save_config(config)
    return {"message": "Intervalo atualizado", "interval_seconds": cfg.interval_seconds}